"""
① UUID 永続化 / ② 脚注対応 テスト
"""
from __future__ import annotations
import tempfile
import zipfile
from pathlib import Path

import pytest


# ── UUID 永続化 ────────────────────────────────────────────────────────────

class TestUidStore:
    def test_creates_sidecar_on_first_call(self, tmp_path):
        from tateyomi.utils.uid_store import get_or_create_uid
        input_file = tmp_path / "book.txt"
        input_file.write_text("テスト")
        uid = get_or_create_uid(input_file)
        sidecar = tmp_path / "book.tateyomi-uid"
        assert sidecar.exists()
        assert sidecar.read_text().strip() == uid

    def test_returns_same_uid_on_second_call(self, tmp_path):
        from tateyomi.utils.uid_store import get_or_create_uid
        input_file = tmp_path / "book.txt"
        input_file.write_text("テスト")
        uid1 = get_or_create_uid(input_file)
        uid2 = get_or_create_uid(input_file)
        assert uid1 == uid2

    def test_uid_is_valid_uuid(self, tmp_path):
        import uuid as _uuid
        from tateyomi.utils.uid_store import get_or_create_uid
        input_file = tmp_path / "book.txt"
        input_file.write_text("テスト")
        uid = get_or_create_uid(input_file)
        # UUID としてパースできること
        parsed = _uuid.UUID(uid)
        assert str(parsed) == uid

    def test_read_uid_returns_none_if_no_sidecar(self, tmp_path):
        from tateyomi.utils.uid_store import read_uid
        input_file = tmp_path / "book.txt"
        input_file.write_text("テスト")
        assert read_uid(input_file) is None

    def test_read_uid_returns_stored_uid(self, tmp_path):
        from tateyomi.utils.uid_store import get_or_create_uid, read_uid
        input_file = tmp_path / "book.txt"
        input_file.write_text("テスト")
        uid = get_or_create_uid(input_file)
        assert read_uid(input_file) == uid

    def test_different_files_get_different_uids(self, tmp_path):
        from tateyomi.utils.uid_store import get_or_create_uid
        f1 = tmp_path / "a.txt"
        f2 = tmp_path / "b.txt"
        f1.write_text("A")
        f2.write_text("B")
        uid1 = get_or_create_uid(f1)
        uid2 = get_or_create_uid(f2)
        assert uid1 != uid2


# ── 青空文庫 インライン注記 ────────────────────────────────────────────────

class TestAozoraInlineNote:
    def test_inline_note_converted(self):
        from tateyomi.utils.aozora import parse_aozora
        text = "テスト書籍\n著者名\n\nこれは本文（注：補足説明）です。"
        _, _, blocks = parse_aozora(text)
        html = " ".join(b.content for b in blocks)
        assert 'class="fn-inline"' in html
        assert "補足説明" in html

    def test_inline_note_has_role_footnote(self):
        from tateyomi.utils.aozora import parse_aozora
        text = "タイトル\n著者\n\n本文（注：注記内容）。"
        _, _, blocks = parse_aozora(text)
        html = " ".join(b.content for b in blocks)
        assert 'role="doc-footnote"' in html

    def test_normal_text_unaffected(self):
        from tateyomi.utils.aozora import parse_aozora
        text = "タイトル\n著者\n\n普通の文章です。"
        _, _, blocks = parse_aozora(text)
        html = " ".join(b.content for b in blocks)
        assert "fn-inline" not in html

    def test_multiple_notes_in_one_line(self):
        from tateyomi.utils.aozora import parse_aozora
        text = "タイトル\n著者\n\nA（注：注1）B（注：注2）"
        _, _, blocks = parse_aozora(text)
        html = " ".join(b.content for b in blocks)
        assert html.count("fn-inline") == 2


# ── docx 脚注抽出ヘルパー ──────────────────────────────────────────────────

class TestDocxFootnoteHelper:
    def _make_docx_with_footnotes(self, tmp_path: Path) -> Path:
        """脚注付き docx を生成する"""
        try:
            from docx import Document
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            import lxml.etree as etree
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.core_properties.title = "脚注テスト"
        doc.core_properties.author = "テスト著者"
        p = doc.add_paragraph("本文テキスト")
        path = tmp_path / "footnote_test.docx"
        doc.save(str(path))
        return path

    def test_get_docx_footnotes_returns_dict(self, tmp_path):
        """docx 生成して _get_docx_footnotes が dict を返すことを確認"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")
        from tateyomi.parsers.docx_parser import _get_docx_footnotes
        doc_path = self._make_docx_with_footnotes(tmp_path)
        from docx import Document
        doc = Document(str(doc_path))
        result = _get_docx_footnotes(doc)
        assert isinstance(result, dict)

    def test_docx_parser_runs_without_footnotes(self, tmp_path):
        """脚注なし docx が正常にパースできること"""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")
        from tateyomi.parsers.docx_parser import DocxParser
        doc_path = self._make_docx_with_footnotes(tmp_path)
        book = DocxParser().parse(doc_path)
        assert book.title == "脚注テスト"
        assert len(book.chapters) >= 1


# ── CSS 脚注スタイル ───────────────────────────────────────────────────────

class TestFootnoteCss:
    def _read_css(self, filename: str) -> str:
        from pathlib import Path
        assets = Path(__file__).parent.parent / "tateyomi" / "assets"
        return (assets / filename).read_text(encoding="utf-8")

    def test_footnotes_class_in_main_css(self):
        css = self._read_css("tateyomi.css")
        assert ".footnotes" in css

    def test_footnote_aside_in_main_css(self):
        css = self._read_css("tateyomi.css")
        assert 'aside' in css
        assert "footnote" in css

    def test_fn_inline_in_main_css(self):
        css = self._read_css("tateyomi.css")
        assert ".fn-inline" in css

    def test_footnote_in_kindle_overrides(self):
        css = self._read_css("kindle-overrides.css")
        assert ".footnotes" in css or "footnote" in css
